import os, sys
import json
import requests
from django.http import JsonResponse
from datetime import datetime
import openai
from django.conf import settings
from whatsapp.models import ChatbotData
from whatsapp.models import UserData
from datetime import datetime
import docx
import re


import environ
env = environ.Env()
env_path = os.path.join(settings.BASE_DIR, '.env')
environ.Env.read_env(env_file=env_path)


def get_media_url(media_id):
    print("Inside get_media_url function: ", media_id)
    whatsapp_token = env('whatsapp_token')
    try:
        url = "https://graph.facebook.com/v12.0/" + str(media_id)
        # payload = {
        #     "messaging_product": "whatsapp",
        #     "to": from_number,
        #     "text": { "body": msg_body }
        # }
        headers = {
            "Authorization": f"Bearer {str(whatsapp_token)}"
        }


        response = requests.get(url, headers=headers)
        print(response.json())
        if response.status_code == 200:
            print("URL retrieved successfully!")
            response_out = response.json()
            print('URL: ', response_out['url'])
            return response_out['url']
        else:
            print("Error retrieved message:", response.text)
    except Exception as _e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print("ERROR: ", exc_type, fname, exc_tb.tb_lineno)
        print("EXCEPTION: ", _e)

def manage_mime_type(response, media_info):
    mime_type = media_info['mime_type']
    doc_name = media_info['doc_name']
    doc_name = doc_name[:5]
    # Get the current time
    current_time = datetime.now()

    # Format the time as 24_11_23_HH_MM_SS.txt
    time_format = current_time.strftime('%y_%m_%d_%H_%M_%S')
    if str(mime_type) == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        print("IT'S A DOCX FILE....")
        file_path = f"/home/arjunmdr/whatsapp_webhook/files/{doc_name}_{time_format}.docx"  # Provide the desired file path
        with open (file_path, 'wb') as file:
            file.write(response.content)
        categorize_media(file_path, media_info)
    elif str(mime_type) == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
        print("IT'S A XLSX FILE....")
        file_path = f"/home/arjunmdr/whatsapp_webhook/files/{doc_name}_{time_format}.xlsx"  # Provide the desired file path
        with open (file_path, 'wb') as file:
            file.write(response.content)
    elif str(mime_type) == 'application/pdf':
        print("IT'S A PDF FILE....")
        file_path = f"/home/arjunmdr/whatsapp_webhook/files/{doc_name}_{time_format}.pdf"  # Provide the desired file path
        with open (file_path, 'wb') as file:
            file.write(response.content)
    else:
        print("NOT SUPPORTED FORMAT FILE....")


def categorize_media(file_path, media_info):
    mime_type = media_info['mime_type']
    # List of keywords to match
    keywords = ["python", "django", "devops", "cicd", "linux", "react", "reactjs", "angular", "Ansible", "Celenium", "docker"]

    if str(mime_type) == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        # Load the docx file
        doc = docx.Document(file_path)
        # Extract the text from the docx file
        text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
        # Check for matching keywords
        matching_keywords = [keyword for keyword in keywords if re.search(r'\b{}\b'.format(keyword), text, re.IGNORECASE)]

        # Print the matching keywords
        print("matching_keywords: ", matching_keywords)
        for keyword in matching_keywords:
            print("Matching keyword:", keyword)

    # elif str(mime_type) == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
    #     print("IT'S A XLSX FILE....")

    # elif str(mime_type) == 'application/pdf':
    #     print("IT'S A PDF FILE....")

    else:
        print("NOT SUPPORTED FORMAT FILE....")

def download_media(url, media_info):
    #url_without_https = url.replace("https://", "")
    whatsapp_token = env('whatsapp_token')
    try:
        headers = {
            "Authorization": f"Bearer {str(whatsapp_token)}"
        }
        print('url: ', url)
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            print("URL download successfully!")
            manage_mime_type(response, media_info)
        else:
            print("Error download message:", response.text)
    except Exception as _e:
        exc_type, exc_obj, exc_tb = sys.exc_info()
        fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
        print("ERROR: ", exc_type, fname, exc_tb.tb_lineno)
        print("EXCEPTION: ", _e)

def manageDocument(body_obj):
    """
    check document messages
    """
    whatsapp_token = env('whatsapp_token')
    phone_number_id = body_obj['entry'][0]['changes'][0]['value']['metadata']['phone_number_id']
    from_number = body_obj['entry'][0]['changes'][0]['value']['messages'][0]['from']
    doc_name = body_obj['entry'][0]['changes'][0]['value']['messages'][0]['document']['filename']
    doc_id = body_obj['entry'][0]['changes'][0]['value']['messages'][0]['document']['id']
    mime_type = body_obj['entry'][0]['changes'][0]['value']['messages'][0]['document']['mime_type']
    media_info = {
        'phone_number_id': phone_number_id,
        'from_number': from_number,
        'doc_name': doc_name,
        'doc_id': doc_id,
        'mime_type': mime_type
    }
    print("phone_number_id: ", phone_number_id)
    print("from_number: ", from_number)
    print("doc_name: ", doc_name)
    print("doc_id: ", doc_id)
    print("mime_type: ", mime_type)
    
    media_url = get_media_url(str(doc_id))
    download_media(str(media_url), media_info)
